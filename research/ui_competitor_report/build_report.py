from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

ROOT = Path('/Users/julia_chen/Documents/ChatGPT/贵客松')
ASSETS = ROOT / 'research/ui_competitor_report/assets'
OUT = ROOT / '灵枢数字员工UI竞品研究与设计建议.docx'

NAVY = '14213D'; BLUE = '3157D5'; CYAN = '20A4A8'; ORANGE = 'F47C20'
LIGHT = 'F3F6FA'; MID = 'D9E1EC'; DARK = '243043'; MUTED = '687386'; WHITE = 'FFFFFF'

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = tcPr.find(qn('w:shd'))
    if shd is None: shd = OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn('w:'+m))
        if node is None: node=OxmlElement('w:'+m); tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr(); el = OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def set_col_width(cell, inches):
    cell.width = Inches(inches); tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'), str(int(inches*1440))); tcW.set(qn('w:type'),'dxa')

def font(run, size=10.5, bold=False, color=DARK, name='Noto Sans CJK SC'):
    run.font.name=name; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    run.font.size=Pt(size); run.bold=bold; run.font.color.rgb=RGBColor.from_string(color)

def para(doc, text='', size=10.5, bold=False, color=DARK, before=0, after=6, align=None):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.15
    if align is not None: p.alignment=align
    font(p.add_run(text), size, bold, color); return p

def bullet(doc, text, color=DARK):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(4); p.paragraph_format.line_spacing=1.15
    font(p.add_run(text),10.2,False,color); return p

def heading(doc, text, level=1):
    p=doc.add_paragraph(style=f'Heading {level}'); p.add_run(text); return p

def callout(doc, label, text, color=BLUE):
    t=doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    c=t.cell(0,0); set_col_width(c,6.5); set_cell_shading(c,LIGHT); set_cell_margins(c,140,170,140,170)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(label+'  '),10.5,True,color); font(p.add_run(text),10.5,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)

def add_image(doc, filename, caption, source):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(3)
    p.add_run().add_picture(str(ASSETS/filename), width=Inches(6.4))
    cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER; cp.paragraph_format.space_after=Pt(2)
    font(cp.add_run(caption),9,True,DARK)
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER; sp.paragraph_format.space_after=Pt(8)
    font(sp.add_run('来源：'+source+'（2026-08-29访问；官方产品页截图）'),8,False,MUTED)

def add_table(doc, headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; set_col_width(c,widths[i]); set_cell_shading(c,NAVY); set_cell_margins(c)
        c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        font(p.add_run(h),9.2,True,WHITE)
    set_repeat_table_header(t.rows[0])
    for ridx,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            c=cells[i]; set_col_width(c,widths[i]); set_cell_margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx%2: set_cell_shading(c,'F7F9FC')
            p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); font(p.add_run(str(val)),8.8,False,DARK)
    doc.add_paragraph().paragraph_format.space_after=Pt(2); return t

doc=Document(); sec=doc.sections[0]
sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=Inches(0.75); sec.left_margin=sec.right_margin=Inches(1)
sec.header_distance=Inches(.35); sec.footer_distance=Inches(.35)

styles=doc.styles
normal=styles['Normal']; normal.font.name='Noto Sans CJK SC'; normal._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC'); normal.font.size=Pt(10.5); normal.font.color.rgb=RGBColor.from_string(DARK)
normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.15
for lvl,size,bef,aft,color in [(1,17,16,8,NAVY),(2,13.5,12,6,BLUE),(3,11.5,8,4,CYAN)]:
    s=styles[f'Heading {lvl}']; s.font.name='Noto Sans CJK SC'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC'); s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft); s.paragraph_format.keep_with_next=True
lb=styles['List Bullet']; lb.font.name='Noto Sans CJK SC'; lb._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC'); lb.font.size=Pt(10.2); lb.paragraph_format.left_indent=Inches(.32); lb.paragraph_format.first_line_indent=Inches(-.17)

header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(header.add_run('灵枢 Global Growth OS · 产品研究'),8.5,False,MUTED)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(footer.add_run('内部讨论稿 · 2026-08-29'),8,False,MUTED)

# Cover
para(doc,'PRODUCT UI RESEARCH',10,True,ORANGE,after=22)
para(doc,'灵枢数字员工 UI',29,True,NAVY,after=3)
para(doc,'竞品研究与产品界面设计建议',21,True,BLUE,after=18)
para(doc,'面向贵州特色产业 B2B 出口的「内容创作 40% · 信号提取 30% · 客服承接 30%」数字员工',12,False,DARK,after=24)
callout(doc,'核心结论','灵枢不应设计成“内容工具＋数据看板＋客服后台”的拼盘，而应设计成一个可被管理、授权、观察和接管的海外内容增长数字员工。')
heading(doc,'研究范围',1)
bullet(doc,'重点样本：Omneky、Creatify、Hermoso、Adanamous、Predis.ai，并参考广告平台原生自动化的交互逻辑。')
bullet(doc,'观察维度：Agent入口、目标设定、创作工作台、数据反馈、行动记录、自主权限、异常审批与商业结果。')
bullet(doc,'截图来自公开产品页面。由于多数后台需要注册，本报告分析的是公开可验证的产品结构与交互表达，不把营销演示图当作完整实装证明。')

heading(doc,'一、结论先行：应该采用“员工驾驶舱”，不是“功能后台”',1)
para(doc,'传统营销软件把页面按功能拆分：素材、脚本、日历、发布、数据、会话。数字员工产品应把页面按管理动作组织：设目标、看它在做什么、审批高风险动作、理解它学到了什么、接管高价值机会。')
add_table(doc,['界面问题','传统工具答案','灵枢应给出的答案'],[
('今天要做什么？','用户自己进入各模块操作','数字员工展示今日计划、进行中任务与预计结果'),
('为什么这样做？','数据散落在多个报表','每个行动附带信号、判断、置信度和目标'),
('什么需要人工？','所有步骤都需要人协调','只把审批、异常和高价值客户放入人工队列'),
('如何衡量？','内容数、播放量、回复量','目标买家触达、有效信号、合格商机与可归因结果'),
], [1.25,2.25,3.0])

heading(doc,'二、竞品界面观察',1)
heading(doc,'2.1 Creatify：用“一个岗位”替代复杂广告菜单',2)
add_image(doc,'creatify-media-buyer.png','图 1  Creatify 直接把产品包装成 AI 效果营销专家','https://creatify.ai/features/media-buyer')
bullet(doc,'优点：第一屏只讲角色、目标和覆盖平台，减少用户理解成本。')
bullet(doc,'可借鉴：灵枢首页应先出现“海外内容增长员工”，而不是“AI创作、信号中心、客服”等模块名称。')
bullet(doc,'不应照搬：Creatify仍以广告账户和效果投放为中心，缺少贵州B2B出口所需的产品证据、渠道商信号和人工业务接管。')

heading(doc,'2.2 Hermoso：Agent在中央，营销工具退到执行层',2)
add_image(doc,'hermoso-home.png','图 2  Hermoso 把自身定位为通用 Agent 可调用的营销执行层','https://hermoso.ai/')
add_image(doc,'hermoso-workflow.png','图 3  Hermoso 的工作室仍保留素材、排期、广告研究等工具视图','https://hermoso.ai/')
bullet(doc,'优点：对话/Agent是主入口，工具是Agent的能力，不要求用户逐一协调。')
bullet(doc,'可借鉴：灵枢可以保留素材、日历和会话页，但默认入口必须是员工驾驶舱和任务流。')
bullet(doc,'风险：完全对话式界面不利于批量审核、跨国家比较和管理经营结果，因此不能只做聊天窗口。')

heading(doc,'2.3 Adanamous：把自主运营拆成四个可理解阶段',2)
add_image(doc,'adanamous-home.png','图 4  Adanamous 以“自主媒体买手”建立角色认知','https://adanamous.com/')
add_image(doc,'adanamous-control.png','图 5  市场机会、创意、执行、优化四段式能力表达','https://adanamous.com/')
bullet(doc,'优点：用“寻找机会—转成创意—上线执行—放大有效结果”讲闭环，而不是罗列功能。')
bullet(doc,'可借鉴：灵枢的业务闭环可改为“发现市场信号—生产内容—客户互动—学习与再行动”。')
bullet(doc,'关键补充：B2B出口必须增加证据与审批层，不能让Agent未经授权承诺价格、认证、交期或渠道政策。')

heading(doc,'2.4 Predis.ai：自动发帖的流程清楚，但仍停留在效率工具',2)
add_image(doc,'creatify-workflow.png','图 6  Predis.ai 的“连接品牌—生成内容—自动发布”三步结构','https://predis.ai/auto-post/')
bullet(doc,'优点：上手路径短、指标直观，适合表现“设一次，持续运行”。')
bullet(doc,'不足：强调节省工时和发帖量，没有呈现为什么创作、从市场学到了什么、是否触达正确买家。')
bullet(doc,'对灵枢的启示：效率指标只放二级；一级必须是买家信号、市场假设和商业机会。')

heading(doc,'2.5 Omneky：Agent连接营销栈，但高风险动作保留确认',2)
add_image(doc,'omneky-mcp.png','图 7  Omneky 把 Agent 与品牌、商品、广告和效果数据连接','https://www.omneky.com/mcp')
bullet(doc,'优点：通过自然语言分析、生成、发布和优化，同时把外部平台连接明确展示。')
bullet(doc,'可借鉴：对花钱、外发和业务承诺设明确确认点；低风险动作自动执行。')
bullet(doc,'界面要求：每个动作必须展示“将影响什么、使用多少预算、基于什么判断、是否可撤回”。')

heading(doc,'三、灵枢推荐的信息架构',1)
callout(doc,'设计原则','一级导航围绕管理数字员工，二级工作台才围绕内容、信号和会话。')
add_table(doc,['一级页面','回答的问题','主要内容'],[
('1. 员工驾驶舱','它负责什么、现在怎么样？','目标、市场、进行中任务、关键结果、异常、待人工事项'),
('2. 内容工作台（40%）','它在生产和发布什么？','策略、选题、素材、脚本、版本、审核、日历、投流实验'),
('3. 信号中心（30%）','海外市场在反馈什么？','品牌、需求、渠道、采购、竞品、风险信号及证据'),
('4. 客户承接（30%）','哪些互动需要回复或转交？','统一收件箱、身份判断、建议回复、线索评分、人工接管'),
('5. 结果与学习','它带来了什么并学到了什么？','买家触达、有效信号、合格商机、内容归因、策略记忆'),
('6. 权限与连接','它能自主做到哪一步？','账号、预算、自动发布、外发、合规红线、审批和审计'),
], [1.3,2.0,3.2])

heading(doc,'四、六个关键页面的界面草案',1)
heading(doc,'4.1 员工驾驶舱：首页不是数据大屏，而是管理面板',2)
add_table(doc,['区域','建议组件','优先级'],[
('顶部目标条','目标国家、买家类型、周期目标、预算、自主等级','P0'),
('今日行动流','正在研究/生成/发布/回复/等待审批的任务时间线','P0'),
('四项核心结果','目标买家触达、有效信号、合格对话、人工接管','P0'),
('Agent判断','本周学到的三件事、下一步建议、置信度和证据','P0'),
('人工收件箱','待审批内容、风险回复、高价值客户、资料缺口','P0'),
('经营趋势','按国家、产品、买家角色查看漏斗和内容贡献','P1'),
], [1.35,3.9,1.25])
para(doc,'推荐布局：左侧窄导航；中央70%为目标与行动时间线；右侧30%为“需要你处理”与风险。不要首屏堆十几个图表。',10.5,True,BLUE)

heading(doc,'4.2 内容工作台：以“内容任务”而不是“编辑器”组织',2)
bullet(doc,'左栏：市场、买家角色、采购阶段、内容目标和证据来源。')
bullet(doc,'中栏：脚本/图文/视频的多版本画布，显示Agent为何选这个角度。')
bullet(doc,'右栏：品牌事实、合规、平台格式、CTA和发布建议。')
bullet(doc,'底部：版本对比、审批、发布/投流、回滚与行动审计。')
callout(doc,'关键差异','每条内容必须绑定“面向谁、解决什么采购问题、使用什么证据、希望触发什么信号”，否则只是普通AIGC。')

heading(doc,'4.3 信号中心：不要做成社媒舆情大屏',2)
bullet(doc,'默认展示“值得行动的信号流”，不是所有点赞评论。')
bullet(doc,'每条信号包含：来源、国家、客户/公司、信号类型、强度、证据、建议动作。')
bullet(doc,'信号类型建议固定为：品牌兴趣、产品需求、采购意向、渠道合作、专业技术、竞品机会、风险。')
bullet(doc,'一键触发：生成解释内容、补充FAQ、提高投流、降低频率、创建客服任务、转交人工。')

heading(doc,'4.4 客户承接：轻客服，强识别与转交',2)
bullet(doc,'三栏式：对话列表—当前对话—客户/公司与信号侧栏。')
bullet(doc,'Agent先完成语言识别、公司识别、买家类型判断和标准资料调用。')
bullet(doc,'界面重点显示“为什么判断为渠道商/采购商”“缺什么资格信息”“是否需要人工”。')
bullet(doc,'不在P0建设工单、售后、物流、退款和完整CRM。')

heading(doc,'4.5 自主权限中心：数字员工必须可控',2)
add_table(doc,['动作','默认模式','建议边界'],[
('研究、分析、生成草稿','自动','保留来源与模型判断'),
('已批准模板的自然发布','自动/批量授权','限定账号、国家、频率'),
('首次发布或敏感行业内容','审批','事实、合规、品牌三项检查'),
('小额投流调整','阈值内自动','日预算、单次增幅、止损线'),
('标准资料回复','自动','只调用已批准知识'),
('价格、交期、认证、独家代理','人工','禁止Agent自主承诺'),
], [2.0,1.55,2.95])

heading(doc,'4.6 结果与学习：展示“它学会了什么”',2)
bullet(doc,'内容层：哪些角度、格式、语言和平台有效。')
bullet(doc,'信号层：哪些国家、买家角色和问题正在升温。')
bullet(doc,'客服层：哪些问题最常见，哪些对话转为有效商机。')
bullet(doc,'组织层：Agent采取了哪些行动、节省多少人工、产生哪些需要业务员接管的机会。')

heading(doc,'五、P0开发范围与4:3:3落地',1)
add_table(doc,['能力包','开发占比','P0必须完成','暂缓'],[
('内容创作','40%','知识约束、策略/选题、多语内容、视频/图文、审核、发布、基础投流','复杂视频模型管理、重型素材资产管理'),
('信号提取','30%','评论/私信/数据接入、七类信号、证据、强度、建议动作、策略回流','全网舆情、过度复杂预测模型'),
('客服承接','30%','统一会话、标准回复、公司/角色识别、资格补全、人工接管','售后工单、退款物流、完整CRM、合同与报价'),
], [1.15,0.85,2.8,1.7])
callout(doc,'P0演示主线','设定一个贵州茶/农特产品的东南亚渠道增长目标 → Agent生成并发布内容 → 捕捉“认证/MOQ/代理”信号 → 自动回复并识别海外渠道商 → 人工接管 → Agent根据反馈调整下一轮内容。')

heading(doc,'六、视觉与交互规范建议',1)
bullet(doc,'视觉气质：可信、专业、克制，避免“炫技型AI霓虹大屏”。贵州产业企业更需要可控与可解释。')
bullet(doc,'状态颜色：蓝色=计划/进行中，绿色=已完成，橙色=需审批，红色=风险/越界，灰色=等待外部输入。')
bullet(doc,'每个Agent动作使用统一行动卡：目标、依据、动作、影响、权限、结果、下一步。')
bullet(doc,'重要AI判断必须展示证据与置信度，允许用户纠正并形成组织记忆。')
bullet(doc,'对话不是唯一入口；批量审核、信号队列和结果对比必须采用结构化界面。')
bullet(doc,'移动端只做审批、异常和高价值客户提醒；完整创作与分析放桌面端。')

heading(doc,'七、最终建议',1)
para(doc,'灵枢前端应当让客户感觉自己“雇佣并管理了一名海外内容增长员工”，而不是购买了一套功能更多的内容平台。',12,True,NAVY,after=10)
add_table(doc,['应强化','应弱化'],[
('员工目标、行动流、判断依据、学习结果','首页堆叠传统流量图表'),
('内容—信号—客服三者的闭环','三个独立系统式导航'),
('授权边界、审批、异常与审计','宣称完全无人监管'),
('目标买家、渠道信号和合格商机','只强调发帖量、播放量、节省工时'),
('企业证据和B2B采购语境','通用消费品爆款模板'),
], [3.25,3.25])
para(doc,'建议下一步直接制作六页高保真原型：员工驾驶舱、内容任务、信号中心、客户承接、自主权限、结果与学习。先用一个真实贵州品牌和一个目标国家跑通演示数据，再扩展多品牌控制塔。')

heading(doc,'附录：公开来源',1)
for s in [
('Omneky MCP','https://www.omneky.com/mcp'),('Creatify AI Media Buyer','https://creatify.ai/features/media-buyer'),('Hermoso','https://hermoso.ai/'),('Adanamous','https://adanamous.com/'),('Predis.ai Auto Post','https://predis.ai/auto-post/')]:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); font(p.add_run(s[0]+'：'),9.2,True,DARK); font(p.add_run(s[1]),9.2,False,BLUE)

doc.save(OUT)
print(OUT)
